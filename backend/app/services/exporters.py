from __future__ import annotations

import base64
import csv
import io
import json
from typing import Iterable

from docx import Document
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas

from app.models.entities import VideoJob


class ExportArtifact(dict):
    file_name: str
    media_type: str
    content_base64: str
    text_preview: str


def build_export(job: VideoJob, kind: str) -> dict[str, str]:
    builders = {
        "markdown": lambda: _artifact("notes.md", "text/markdown", render_markdown(job).encode("utf-8"), render_markdown(job)),
        "json": lambda: _artifact("notes.json", "application/json", render_json(job).encode("utf-8"), render_json(job)),
        "txt": lambda: _artifact("transcript.txt", "text/plain", render_txt_transcript(job).encode("utf-8"), render_txt_transcript(job)),
        "srt": lambda: _artifact("transcript.srt", "application/x-subrip", render_srt(job).encode("utf-8"), render_srt(job)),
        "vtt": lambda: _artifact("transcript.vtt", "text/vtt", render_vtt(job).encode("utf-8"), render_vtt(job)),
        "csv": lambda: _artifact("action-items.csv", "text/csv", render_action_items_csv(job).encode("utf-8"), render_action_items_csv(job)),
        "pdf": lambda: _artifact("notes.pdf", "application/pdf", render_pdf(job), render_markdown(job)),
        "docx": lambda: _artifact(
            "notes.docx",
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            render_docx(job),
            render_markdown(job),
        ),
    }
    return builders[kind]()


def render_markdown(job: VideoJob) -> str:
    summary_sections = "\n\n".join(f"## {section.section}\n\n{section.content_markdown}" for section in job.summaries)
    action_items = "\n".join(
        f"- [{'x' if item.completed else ' '}] {item.description} ({item.owner or 'Unassigned'})"
        for item in job.action_items
    )
    timestamps = "\n".join(
        f"- {stamp.label} ({_clock(stamp.start_seconds)}): {stamp.description}" for stamp in job.timestamps
    )
    return (
        f"# {job.title}\n\n"
        f"Mode: {job.mode.value}\n\n"
        f"## Overview\n\n{job.metadata_json.get('summary_overview', 'Structured notes generated from transcript evidence.')}\n\n"
        f"{summary_sections}\n\n"
        f"## Action Items\n\n{action_items or '- None'}\n\n"
        f"## Important Timestamps\n\n{timestamps or '- None'}"
    )


def render_json(job: VideoJob) -> str:
    payload = {
        "id": job.id,
        "title": job.title,
        "mode": job.mode.value,
        "status": job.status.value,
        "source_type": job.source_type.value,
        "source_url": job.source_url,
        "duration_seconds": job.duration_seconds,
        "metadata": job.metadata_json,
        "summaries": [summary.structured_json for summary in job.summaries],
        "action_items": [
            {
                "description": item.description,
                "owner": item.owner,
                "due_hint": item.due_hint,
                "completed": item.completed,
            }
            for item in job.action_items
        ],
        "timestamps": [
            {
                "label": stamp.label,
                "description": stamp.description,
                "start_seconds": stamp.start_seconds,
                "end_seconds": stamp.end_seconds,
            }
            for stamp in job.timestamps
        ],
    }
    return json.dumps(payload, indent=2)


def render_txt_transcript(job: VideoJob) -> str:
    return "\n".join(
        f"[{_clock(chunk.start_seconds)} - {_clock(chunk.end_seconds)}] {chunk.speaker or 'Speaker'}: {chunk.text}"
        for chunk in job.transcript_chunks
    )


def render_srt(job: VideoJob) -> str:
    lines: list[str] = []
    for index, chunk in enumerate(job.transcript_chunks, start=1):
        start = _format_time(chunk.start_seconds, separator=",")
        end = _format_time(chunk.end_seconds, separator=",")
        lines.extend([str(index), f"{start} --> {end}", f"{chunk.speaker or 'Speaker'}: {chunk.text}", ""])
    return "\n".join(lines)


def render_vtt(job: VideoJob) -> str:
    lines = ["WEBVTT", ""]
    for chunk in job.transcript_chunks:
        start = _format_time(chunk.start_seconds)
        end = _format_time(chunk.end_seconds)
        lines.extend([f"{start} --> {end}", f"{chunk.speaker or 'Speaker'}: {chunk.text}", ""])
    return "\n".join(lines)


def render_action_items_csv(job: VideoJob) -> str:
    output = io.StringIO()
    writer = csv.writer(output)
    writer.writerow(["description", "owner", "due_hint", "completed"])
    for item in job.action_items:
        writer.writerow([item.description, item.owner or "", item.due_hint or "", item.completed])
    return output.getvalue()


def render_pdf(job: VideoJob) -> bytes:
    buffer = io.BytesIO()
    pdf = canvas.Canvas(buffer, pagesize=A4)
    width, height = A4
    y = height - 48
    lines = _wrapped_lines(
        [job.title, "", *render_markdown(job).splitlines()],
        max_chars=92,
    )
    pdf.setTitle(job.title)
    for line in lines:
        if y < 48:
            pdf.showPage()
            y = height - 48
        pdf.drawString(40, y, line)
        y -= 14
    pdf.save()
    return buffer.getvalue()


def render_docx(job: VideoJob) -> bytes:
    document = Document()
    document.add_heading(job.title, level=1)
    document.add_paragraph(f"Mode: {job.mode.value}")
    for section in job.summaries:
        document.add_heading(section.section, level=2)
        document.add_paragraph(section.content_markdown)
    document.add_heading("Action Items", level=2)
    for item in job.action_items:
        document.add_paragraph(
            f"{item.description} | Owner: {item.owner or 'Unassigned'} | Due: {item.due_hint or 'None'}",
            style="List Bullet",
        )
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _artifact(file_name: str, media_type: str, content: bytes, preview: str) -> dict[str, str]:
    return {
        "file_name": file_name,
        "media_type": media_type,
        "content_base64": base64.b64encode(content).decode("utf-8"),
        "text_preview": preview[:1200],
    }


def _wrapped_lines(lines: Iterable[str], *, max_chars: int) -> list[str]:
    wrapped: list[str] = []
    for line in lines:
        current = line.strip()
        if not current:
            wrapped.append("")
            continue
        while len(current) > max_chars:
            split_at = current.rfind(" ", 0, max_chars)
            if split_at <= 0:
                split_at = max_chars
            wrapped.append(current[:split_at])
            current = current[split_at:].strip()
        wrapped.append(current)
    return wrapped


def _clock(seconds: float) -> str:
    whole = int(seconds)
    hours = whole // 3600
    minutes = (whole % 3600) // 60
    secs = whole % 60
    return f"{hours:02}:{minutes:02}:{secs:02}"


def _format_time(seconds: float, *, separator: str = ".") -> str:
    whole = int(seconds)
    hours = whole // 3600
    minutes = (whole % 3600) // 60
    secs = whole % 60
    millis = max(int(round((seconds - whole) * 1000)), 0)
    return f"{hours:02}:{minutes:02}:{secs:02}{separator}{millis:03}"
